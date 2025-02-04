from docxtpl import DocxTemplate
import tkinter as tk
from tkinter import messagebox, filedialog
from tkcalendar import DateEntry
from docx import Document
import sys
import os
import json  # 新增，用于处理JSON数据

# 定义缓存文件的路径
CACHE_FILE = 'cache.json'

# 定义一个函数，填充文档
def fill_document(data):
    # 获取模板文档的路径
    if hasattr(sys, '_MEIPASS'):
        template_path = os.path.join(sys._MEIPASS, '模板文档.docx')  # 修改路径
    else:
        template_path = '模板文档.docx'
    
    # 加载模板文档
    doc = DocxTemplate(template_path)
    # 渲染模板，填充数据
    doc.render(data)
    # 保存临时文档
    temp_doc_path = '临时文档.docx'
    doc.save(temp_doc_path)
    # 打开临时文档
    document = Document(temp_doc_path)

    # 遍历文档中的段落，替换未填充的“目标1”占位符为“以下空白”
    for paragraph in document.paragraphs:
        if '{{目标1}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{目标1}}', '以下空白')

    # 遍历文档中的所有表格，替换未填充的“目标1”占位符为“以下空白”
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{目标1}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{目标1}}', '以下空白')

    # 返回处理后的文档对象
    return document

# 定义保存缓存的函数
def save_cache():
    data = get_form_data()
    # 移除不需要缓存的字段
    fields_to_exclude = [
        '表格编号', '填表日期', '号1', '号2', '号3', '号4', '号5', '号6', '号7',
        '姓名1', '姓名2', '姓名3', '姓名4', '姓名5', '姓名6', '姓名7',
        '号码1', '号码2', '号码3', '号码4', '号码5', '号码6', '号码7',
        '对象', '目标', '目标1', '目标2', '目标3', '目标4', '目标5', '目标6', '目标7'
    ]
    for field in fields_to_exclude:
        if field in data:
            del data[field]
    try:
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        # print("缓存已保存")  # 移除调试输出
    except Exception as e:
        print(f"保存缓存时出错：{e}")

# 定义加载缓存的函数
def load_cache():
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # print("缓存已加载")  # 移除调试输出
            return data
        except Exception as e:
            print(f"加载缓存时出错：{e}")
    return {}

# 定义获取表单数据的函数
def get_form_data():
    def get_goals(goal_vars, goal_entries):
        return ';'.join([f"{var.get()}:{entry.get()}" for var, entry in zip(goal_vars, goal_entries) if entry.get().strip()])

    def count_goals(*goal_lists):
        return sum(len(goals.split(';')) for goals in goal_lists if goals)

    def get_name(entry, index):
        name = entry.get().strip()
        if name == f"如没姓名，请填不详{index}":
            return ""
        return name

    def get_number(entry):
        number = entry.get().strip()
        if number == "如没号码，请填不详":
            return ""
        return number

    data = {
        '表格编号': entry_table_number.get(),
        '填表单位': variable_unit.get(),
        '填表日期': entry_date.get(),
        '案事件名称': entry_case_name.get(),
        '案事件类别': variable_case_category.get(),
        '文书号': entry_document_number.get(),
        '承办人': entry_person.get(),
        '联系方式': entry_contact.get(),
        '基本情况': text_basic_info.get("1.0", tk.END).strip(),
        '对象': entry_dui.get(),
        '目标': '',  # 目标将在后面自动填充
        '号1': entry_number_label1.get(),
        '姓名1': get_name(entry_name1, 1),
        '号码1': get_number(entry_number1),
        '目标1': get_goals(goal_vars1, entry_goals1),
        '号2': entry_number_label2.get(),
        '姓名2': get_name(entry_name2, 2),
        '号码2': get_number(entry_number2),
        '目标2': get_goals(goal_vars2, entry_goals2),
        '号3': entry_number_label3.get(),
        '姓名3': get_name(entry_name3, 3),
        '号码3': get_number(entry_number3),
        '目标3': get_goals(goal_vars3, entry_goals3),
        '号4': entry_number_label4.get(),
        '姓名4': get_name(entry_name4, 4),
        '号码4': get_number(entry_number4),
        '目标4': get_goals(goal_vars4, entry_goals4),
        '号5': entry_number_label5.get(),
        '姓名5': get_name(entry_name5, 5),
        '号码5': get_number(entry_number5),
        '目标5': get_goals(goal_vars5, entry_goals5),
        '号6': entry_number_label6.get(),
        '姓名6': get_name(entry_name6, 6),
        '号码6': get_number(entry_number6),
        '目标6': get_goals(goal_vars6, entry_goals6),
        '号7': entry_number_label7.get(),
        '姓名7': get_name(entry_name7, 7),
        '号码7': get_number(entry_number7),
        '目标7': get_goals(goal_vars7, entry_goals7),
    }

    # 自动填充‘目标’
    total_goals = count_goals(
        data['目标1'], data['目标2'], data['目标3'],
        data['目标4'], data['目标5'], data['目标6'], data['目标7']
    )
    data['目标'] = number_to_chinese(str(total_goals))

    return data

# 定义更新总目标的函数
def update_total_goals():
    data = get_form_data()
    total_goals = data['目标']
    entry_mu.delete(0, tk.END)
    entry_mu.insert(0, total_goals)

# 定义更新对象的函数
def update_object():
    count = 0
    for i in range(1, 8):
        number_label_value = globals()[f'entry_number_label{i}'].get().strip()
        if number_label_value:
            count += 1
    entry_dui.delete(0, tk.END)
    entry_dui.insert(0, number_to_chinese(str(count)))

# 定义设置表单数据的函数
def set_form_data(data):
    def set_goals(goal_data, goal_vars, goal_entries):
        goals = goal_data.split(';')
        for i, goal in enumerate(goals):
            if ':' in goal:
                var, entry = goal.split(':', 1)  # 确保只分割一次，避免多余的分号导致错误
                goal_vars[i].set(var)
                goal_entries[i].delete(0, tk.END)
                goal_entries[i].insert(0, entry)

    entry_table_number.delete(0, tk.END)
    entry_table_number.insert(0, data.get('表格编号', ''))

    variable_unit.set(data.get('填表单位', units[0]))

    try:
        date_parts = data.get('填表日期', '').split('年')
        if len(date_parts) >= 2:
            year = int(date_parts[0])
            month_day = date_parts[1].split('月')
            if len(month_day) >= 2:
                month = int(month_day[0])
                day = int(month_day[1].replace('日', ''))
                entry_date.set_date(f"{year}-{month}-{day}")
    except Exception as e:
        print(f"设置日期时出错：{e}")

    entry_case_name.delete(0, tk.END)
    entry_case_name.insert(0, data.get('案事件名称', ''))

    variable_case_category.set(data.get('案事件类别', case_categories[0]))

    entry_document_number.delete(0, tk.END)
    entry_document_number.insert(0, data.get('文书号', ''))

    entry_person.delete(0, tk.END)
    entry_person.insert(0, data.get('承办人', ''))

    entry_contact.delete(0, tk.END)
    entry_contact.insert(0, data.get('联系方式', ''))

    text_basic_info.delete("1.0", tk.END)
    text_basic_info.insert(tk.END, data.get('基本情况', ''))

    entry_dui.delete(0, tk.END)
    entry_dui.insert(0, data.get('对象', ''))

    entry_mu.delete(0, tk.END)
    entry_mu.insert(0, data.get('目标', ''))

    # 设置人员信息
    entry_number_label1.delete(0, tk.END)
    entry_number_label1.insert(0, data.get('号1', ''))
    entry_name1.delete(0, tk.END)
    entry_name1.insert(0, data.get('姓名1', f"如没姓名，请填不详1"))
    entry_name1.config(fg='grey')
    entry_number1.delete(0, tk.END)
    entry_number1.insert(0, data.get('号码1', '如没号码，请填不详'))
    entry_number1.config(fg='grey')
    set_goals(data.get('目标1', ''), goal_vars1, entry_goals1)

    # 按照同样的方式设置其他人员信息
    for i in range(2, 8):
        globals()[f'entry_number_label{i}'].delete(0, tk.END)
        globals()[f'entry_number_label{i}'].insert(0, data.get(f'号{i}', ''))
        globals()[f'entry_name{i}'].delete(0, tk.END)
        globals()[f'entry_name{i}'].insert(0, data.get(f'姓名{i}', f"如没姓名，请填不详{i}"))
        globals()[f'entry_name{i}'].config(fg='grey')
        globals()[f'entry_number{i}'].delete(0, tk.END)
        globals()[f'entry_number{i}'].insert(0, data.get(f'号码{i}', '如没号码，请填不详'))
        globals()[f'entry_number{i}'].config(fg='grey')
        set_goals(data.get(f'目标{i}', ''), globals()[f'goal_vars{i}'], globals()[f'entry_goals{i}'])

# 定义保存文档的函数
def save_document():
    data = get_form_data()

    # 检查必填字段是否为空
    required_fields = [
        '表格编号', '填表单位', '填表日期', '案事件名称', '案事件类别',
        '承办人', '联系方式', '基本情况', '对象', '目标',
        '号1', '姓名1', '号码1', '目标1'
    ]
    for field in required_fields:
        if not data[field]:
            messagebox.showerror("错误", f"{field} 不能为空")
            return

    # 检查目标和姓名的对应关系
    for i in range(1, 8):
        if data[f'目标{i}'] and not data[f'姓名{i}']:
            messagebox.showerror("错误", f"当目标{i}有值时，姓名{i}不能为空")
            return

    # 更新缓存
    save_cache()

    # 处理日期格式
    date = entry_date.get_date()
    data['填表日期'] = f"{date.year}年{date.month}月{date.day}日"

    # 如果目标1为空，将其设置为“以下空白”
    data['目标1'] = data['目标1'] if data['目标1'].strip() else '以下空白'

    # 确保目标2至目标7的第一个空值设置成‘以下空白’
    for i in range(2, 8):
        if not data[f'目标{i}']:
            data[f'目标{i}'] = '以下空白'
            break

    # 确保姓名1至姓名7的第一个空值设置成‘以下空白’
    for i in range(1, 8):
        if not data[f'姓名{i}']:
            data[f'姓名{i}'] = '以下空白'
            break

    # 确保号码1至号码7的第一个空值设置成‘以下空白’
    for i in range(1, 8):
        if not data[f'号码{i}']:
            data[f'号码{i}'] = '以下空白'
            break

    # print(data)  # 移除调试输出
    document = fill_document(data)
    output_path = filedialog.asksaveasfilename(
        title="保存文档",
        defaultextension=".docx",
        filetypes=[("Word 文档", "*.docx")]
    )
    if output_path:
        document.save(output_path)
        messagebox.showinfo("保存成功", f"文档已保存到 {output_path}")
        # 删除临时文档
        if os.path.exists('临时文档.docx'):
            os.remove('临时文档.docx')

# 定义预览文档的函数
def preview_document():
    data = get_form_data()

    # 更新缓存
    save_cache()

    # 处理日期格式
    date = entry_date.get_date()
    data['填表日期'] = f"{date.year}年{date.month}月{date.day}日"

    # 如果目标1为空，将其设置为“以下空白”
    data['目标1'] = data['目标1'] if data['目标1'].strip() else '以下空白'

    # 确保目标2至目标7的第一个空值设置成‘以下空白’
    for i in range(2, 8):
        if not data[f'目标{i}']:
            data[f'目标{i}'] = '以下空白'
            break

    # 确保姓名1至姓名7的第一个空值设置成‘以下空白’
    for i in range(1, 8):
        if not data[f'姓名{i}']:
            data[f'姓名{i}'] = '以下空白'
            break

    # 确保号码1至号码7的第一个空值设置成‘以下空白’
    for i in range(1, 8):
        if not data[f'号码{i}']:
            data[f'号码{i}'] = '以下空白'
            break

    # print(data)  # 移除调试输出
    document = fill_document(data)
    preview_path = '预览文档.docx'
    document.save(preview_path)
    messagebox.showinfo("预览成功", f"预览文档已生成在当前目录：{preview_path}")
    # 删除临时文档
    if os.path.exists('临时文档.docx'):
        os.remove('临时文档.docx')

# 定义打印文档的函数
def print_document():
    data = get_form_data()

    # 更新缓存
    save_cache()

    # 处理日期格式
    date = entry_date.get_date()
    data['填表日期'] = f"{date.year}年{date.month}月{date.day}日"

    # 如果目标1为空，将其设置为“以下空白”
    data['目标1'] = data['目标1'] if data['目标1'].strip() else '以下空白'

    # 确保目标2至目标7的第一个空值设置成‘以下空白’
    for i in range(2, 8):
        if not data[f'目标{i}']:
            data[f'目标{i}'] = '以下空白'
            break

    # 确保姓名1至姓名7的第一个空值设置成‘以下空白’
    for i in range(1, 8):
        if not data[f'姓名{i}']:
            data[f'姓名{i}'] = '以下空白'
            break

    # 确保号码1至号码7的第一个空值设置成‘以下空白’
    for i in range(1, 8):
        if not data[f'号码{i}']:
            data[f'号码{i}'] = '以下空白'
            break

    # print(data)  # 移除调试输出
    document = fill_document(data)
    print_path = '打印文档.docx'
    document.save(print_path)
    
    try:
        os.startfile(print_path, "print")
        messagebox.showinfo("打印成功", f"文档已发送到打印机：{print_path}")
    except Exception as e:
        messagebox.showerror("打印失败", f"打印文档时出错：{e}")

    # 删除临时文档
    if os.path.exists('临时文档.docx'):
        os.remove('临时文档.docx')

# 定义一个函数，将数字转换为大写汉字
def number_to_chinese(num):
    num_dict = {
        '1': '壹',
        '2': '贰',
        '3': '叁',
        '4': '肆',
        '5': '伍',
        '6': '陆',
        '7': '柒',
        '8': '捌',
        '9': '玖',
        '0': '零'
    }
    return ''.join(num_dict.get(digit, digit) for digit in str(num))

# 创建 GUI 界面
root = tk.Tk()
root.title("文档填充工具")

# 定义填表单位的选项
units = [
    "韶关市公安局曲江分局出入境户政科",
    "韶关市公安局曲江分局附城派出所",
    "韶关市公安局曲江分局马坝派出所",
    "韶关市公安局曲江分局白土派出所",
    "韶关市公安局曲江分局指挥中心",
    "韶关市公安局曲江分局禁毒大队",
    "韶关市公安局曲江分局治安大队",
    "韶关市公安局曲江分局经侦大队",
    "韶关市公安局曲江分局国保大队",
    "韶关市公安局曲江分局网警大队",
    "韶关市公安局曲江分局樟市派出所",
    "韶关市公安局曲江分局沙溪派出所",
    "韶关市公安局曲江分局大塘派出所",
    "韶关市公安局曲江分局乌石派出所",
    "韶关市公安局曲江分局小坑派出所",
    "韶关市公安局曲江分局枫湾派出所",
    "韶关市公安局曲江分局罗坑派出所",
    "韶关市公安局曲江分局森林警察大队",
    "韶关市公安局曲江分局交警大队"
]

# 定义案事件类别的选项
case_categories = ["刑事案件", "行政案件"]

# 定义目标的选项
goal_options = ["微信号", "微信ID"]

# 定义变量
variable_unit = tk.StringVar(root)
variable_unit.set(units[0])  # 设置默认选项

# 使用 grid 布局
# 创建主要的框架
main_frame = tk.Frame(root)
main_frame.pack(padx=10, pady=10)

# 表格编号
tk.Label(main_frame, text="表格编号").grid(row=0, column=0, sticky='e', pady=2)
entry_table_number = tk.Entry(main_frame)
entry_table_number.grid(row=0, column=1, sticky='w', pady=2)

# 填表单位
tk.Label(main_frame, text="填表单位").grid(row=0, column=2, sticky='e', pady=2)
optionmenu_unit = tk.OptionMenu(main_frame, variable_unit, *units)
optionmenu_unit.grid(row=0, column=3, sticky='w', pady=2)

# 填表日期
tk.Label(main_frame, text="填表日期").grid(row=0, column=4, sticky='e', pady=2)
entry_date = DateEntry(main_frame, date_pattern='yyyy年m月d日')
entry_date.grid(row=0, column=5, sticky='w', pady=2)

# 案事件名称
tk.Label(main_frame, text="案事件名称").grid(row=1, column=0, sticky='e', pady=2)
entry_case_name = tk.Entry(main_frame, width=50)
entry_case_name.grid(row=1, column=1, columnspan=5, sticky='w', pady=2)

# 案事件类别
tk.Label(main_frame, text="案事件类别").grid(row=2, column=0, sticky='e', pady=2)
variable_case_category = tk.StringVar(main_frame)
variable_case_category.set(case_categories[0])  # 设置默认选项
optionmenu_case_category = tk.OptionMenu(main_frame, variable_case_category, *case_categories)
optionmenu_case_category.grid(row=2, column=1, columnspan=5, sticky='w', pady=2)

# 文书号
tk.Label(main_frame, text="文书号").grid(row=3, column=0, sticky='e', pady=2)
entry_document_number = tk.Entry(main_frame, width=50)
entry_document_number.grid(row=3, column=1, columnspan=5, sticky='w', pady=2)

# 承办人
tk.Label(main_frame, text="承办人").grid(row=4, column=0, sticky='e', pady=2)
entry_person = tk.Entry(main_frame)
entry_person.grid(row=4, column=1, sticky='w', pady=2)

# 联系方式
tk.Label(main_frame, text="联系方式").grid(row=4, column=2, sticky='e', pady=2)
entry_contact = tk.Entry(main_frame)
entry_contact.grid(row=4, column=3, sticky='w', pady=2)

# 基本情况
tk.Label(main_frame, text="基本情况").grid(row=5, column=0, sticky='ne', pady=2)
text_basic_info = tk.Text(main_frame, height=5, width=50)
text_basic_info.grid(row=5, column=1, columnspan=5, sticky='w', pady=2)

# 定义一个函数，创建号、姓名、号码、目标的输入行
def create_person_row(row, index):
    # 号
    tk.Label(main_frame, text=f"号{index}").grid(row=row, column=0, sticky='e', pady=2)
    number_label_entry = tk.Entry(main_frame, width=10)
    number_label_entry.grid(row=row, column=1, sticky='w', pady=2, padx=0)
    number_label_entry.bind("<KeyRelease>", lambda event: update_object())  # 确保在每次更新‘号’的值时调用

    # 姓名
    tk.Label(main_frame, text=f"姓名{index}").grid(row=row, column=2, sticky='e', pady=2)
    name_entry = tk.Entry(main_frame, width=20, fg='grey')
    name_entry.grid(row=row, column=3, sticky='w', pady=2, padx=0)
    name_entry.insert(0, f"如没姓名，请填不详{index}")  # 增加输入提示

    def on_focus_in(event):
        if name_entry.get() == f"如没姓名，请填不详{index}":
            name_entry.delete(0, tk.END)
            name_entry.config(fg='black')

    def on_focus_out(event):
        if not name_entry.get():
            name_entry.insert(0, f"如没姓名，请填不详{index}")
            name_entry.config(fg='grey')

    name_entry.bind("<FocusIn>", on_focus_in)
    name_entry.bind("<FocusOut>", on_focus_out)

    # 号码
    tk.Label(main_frame, text=f"号码{index}").grid(row=row, column=4, sticky='e', pady=2)
    number_entry = tk.Entry(main_frame, width=30, fg='grey')
    number_entry.grid(row=row, column=5, sticky='w', pady=2, padx=0)
    number_entry.insert(0, "如没号码，请填不详")  # 增加输入提示

    def on_focus_in_number(event):
        if number_entry.get() == "如没号码，请填不详":
            number_entry.delete(0, tk.END)
            number_entry.config(fg='black')

    def on_focus_out_number(event):
        if not number_entry.get():
            number_entry.insert(0, "如没号码，请填不详")
            number_entry.config(fg='grey')

    number_entry.bind("<FocusIn>", on_focus_in_number)
    number_entry.bind("<FocusOut>", on_focus_out_number)

    # 目标
    tk.Label(main_frame, text=f"目标{index}").grid(row=row, column=6, sticky='e', pady=2)
    goal_vars = []
    goal_entries = []
    for i in range(3):  # 假设最多有3个目标
        goal_var = tk.StringVar(main_frame)
        goal_var.set(goal_options[0])  # 设置默认选项
        goal_optionmenu = tk.OptionMenu(main_frame, goal_var, *goal_options)
        goal_optionmenu.grid(row=row, column=7 + i * 2, sticky='w', pady=2, padx=0)
        goal_entry = tk.Entry(main_frame, width=15)
        goal_entry.grid(row=row, column=8 + i * 2, sticky='w', pady=2, padx=2)
        goal_vars.append(goal_var)
        goal_entries.append(goal_entry)

        # 当目标输入框有值时，更新总目标
        goal_entry.bind("<KeyRelease>", lambda event: update_total_goals())

    # 当姓名输入框有值时，填入号的默认序号，并设置对象的值
    def on_name_entry_change(*args):
        if name_entry.get().strip() and name_entry.get() != f"如没姓名，请填不详{index}":
            number_label_entry.delete(0, tk.END)
            number_label_entry.insert(0, str(index))
            update_object()

    name_entry.bind("<KeyRelease>", on_name_entry_change)

    return number_label_entry, name_entry, number_entry, goal_vars, goal_entries

# 创建人员信息输入行
start_row = 6
entry_number_label1, entry_name1, entry_number1, goal_vars1, entry_goals1 = create_person_row(start_row, 1)
entry_number_label2, entry_name2, entry_number2, goal_vars2, entry_goals2 = create_person_row(start_row+1, 2)
entry_number_label3, entry_name3, entry_number3, goal_vars3, entry_goals3 = create_person_row(start_row+2, 3)
entry_number_label4, entry_name4, entry_number4, goal_vars4, entry_goals4 = create_person_row(start_row+3, 4)
entry_number_label5, entry_name5, entry_number5, goal_vars5, entry_goals5 = create_person_row(start_row+4, 5)
entry_number_label6, entry_name6, entry_number6, goal_vars6, entry_goals6 = create_person_row(start_row+5, 6)
entry_number_label7, entry_name7, entry_number7, goal_vars7, entry_goals7 = create_person_row(start_row+6, 7)

# 添加 '对象' 和 '目标' 输入框，并排放置
tk.Label(main_frame, text="对象").grid(row=start_row+7, column=0, sticky='e', pady=2)
entry_dui = tk.Entry(main_frame)
entry_dui.grid(row=start_row+7, column=1, columnspan=2, sticky='w', pady=2)

tk.Label(main_frame, text="目标").grid(row=start_row+7, column=3, sticky='e', pady=2)
entry_mu = tk.Entry(main_frame)
entry_mu.grid(row=start_row+7, column=4, columnspan=2, sticky='w', pady=2)

# 调整列的对齐方式
for i in range(8):
    main_frame.columnconfigure(i, weight=1)

# 创建按钮
frame_buttons = tk.Frame(root)
frame_buttons.pack(padx=10, pady=10)

btn_save = tk.Button(frame_buttons, text="保存文档", command=save_document, width=15)
btn_save.pack(side="left", padx=5)

btn_preview = tk.Button(frame_buttons, text="预览文档", command=preview_document, width=15)
btn_preview.pack(side="left", padx=5)

btn_print = tk.Button(frame_buttons, text="打印文档", command=print_document, width=15)
btn_print.pack(side="left", padx=5)

# 在程序启动时，加载缓存并设置表单数据
cached_data = load_cache()
if (cached_data):
    set_form_data(cached_data)

# 在程序关闭时，保存缓存
def on_closing():
    save_cache()
    root.destroy()

root.protocol("WM_DELETE_WINDOW", on_closing)

root.mainloop()
